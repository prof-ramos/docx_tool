import streamlit as st
import os
import glob
import shutil
import time
import pandas as pd
from docx import Document
from docling.document_converter import DocumentConverter
import unicodedata
import re

# --- Page Configuration ---
st.set_page_config(
    page_title="DOCX Processor Pro",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- Custom CSS ---
st.markdown("""
    <style>
    .main {
        padding-top: 2rem;
    }
    .stButton>button {
        width: 100%;
        border-radius: 8px;
        height: 3em;
        font-weight: bold;
    }
    .success-box {
        padding: 1rem;
        border-radius: 8px;
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        color: #155724;
        margin-bottom: 1rem;
    }
    .error-box {
        padding: 1rem;
        border-radius: 8px;
        background-color: #f8d7da;
        border: 1px solid #f5c6cb;
        color: #721c24;
        margin-bottom: 1rem;
    }
    .stat-card {
        background-color: #f8f9fa;
        border-radius: 8px;
        padding: 1rem;
        text-align: center;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    .stat-value {
        font-size: 2rem;
        font-weight: bold;
        color: #007bff;
    }
    .stat-label {
        color: #6c757d;
        font-size: 0.9rem;
        text-transform: uppercase;
        letter-spacing: 1px;
    }
    </style>
""", unsafe_allow_html=True)

# --- Helper Functions ---
def to_snake_case(name):
    """Converts a string to snake_case."""
    name = unicodedata.normalize('NFKD', name).encode('ASCII', 'ignore').decode('ASCII')
    name = name.lower()
    name = re.sub(r'[^a-z0-9]', '_', name)
    name = re.sub(r'_+', '_', name)
    return name.strip('_')

def remove_colors(input_path, output_path):
    """Removes colors and highlighting from a DOCX file."""
    try:
        doc = Document(input_path)

        def clear_shading(element):
            if element._element.xpath('.//w:shd'):
                for shd in element._element.xpath('.//w:shd'):
                    shd.getparent().remove(shd)

        # Process paragraphs
        for paragraph in doc.paragraphs:
            clear_shading(paragraph)
            for run in paragraph.runs:
                run.font.color.rgb = None
                if run.font.highlight_color is not None:
                    run.font.highlight_color = None

                rPr = run._element.get_or_add_rPr()
                if rPr.xpath('./w:highlight'):
                    for highlight in rPr.xpath('./w:highlight'):
                        highlight.getparent().remove(highlight)
                if rPr.xpath('./w:shd'):
                    for shd in rPr.xpath('./w:shd'):
                        shd.getparent().remove(shd)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    clear_shading(cell)
                    for paragraph in cell.paragraphs:
                        clear_shading(paragraph)
                        for run in paragraph.runs:
                            run.font.color.rgb = None
                            if run.font.highlight_color is not None:
                                run.font.highlight_color = None

                            rPr = run._element.get_or_add_rPr()
                            if rPr.xpath('./w:highlight'):
                                for highlight in rPr.xpath('./w:highlight'):
                                    highlight.getparent().remove(highlight)
                            if rPr.xpath('./w:shd'):
                                for shd in rPr.xpath('./w:shd'):
                                    shd.getparent().remove(shd)

        doc.save(output_path)
        return True, "Colors removed successfully"
    except Exception as e:
        return False, str(e)

def convert_docx_to_md(file_path, output_dir):
    """Converts DOCX to Markdown using Docling."""
    try:
        converter = DocumentConverter()
        result = converter.convert(file_path)
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        output_file = os.path.join(output_dir, f"{base_name}.md")
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(result.document.export_to_markdown())
        return True, output_file
    except Exception as e:
        return False, str(e)

# --- Sidebar Configuration ---
with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/2991/2991112.png", width=64)
    st.title("Settings")

    st.markdown("### 📂 Directories")
    input_dir = st.text_input("Input Directory", value="Administrativo", help="Folder containing .docx files")
    output_dir = st.text_input("Output Directory", value="Output", help="Folder to save processed files")

    st.markdown("### ⚙️ Operations")
    rename_option = st.toggle("Rename to snake_case", value=True, help="Example: 'My File.docx' -> 'my_file.docx'")
    remove_colors_option = st.toggle("Remove Colors/Highlights", value=False, help="Strips all color formatting from text")
    convert_md_option = st.toggle("Convert to Markdown", value=False, help="Generates a corresponding .md file")

    st.markdown("---")
    st.markdown("### 🚀 Actions")
    # Using a form to prevent accidental re-runs
    # with st.form("process_form"):
    #     submitted = st.form_submit_button("Start Processing")
    # Direct button for better feedback loop in this simple app
    start_btn = st.button("Start Processing", type="primary")

# --- Main Content ---
st.title("📄 DOCX Processor Pro")
st.markdown("Batch process your documents with ease. Rename, clean, and convert.")

# Initial State Check
if not os.path.exists(input_dir):
    st.warning(f"⚠️ Input directory `{input_dir}` does not exist. Please check the path.")
    st.stop()

# Scan for files
docx_files = glob.glob(os.path.join(input_dir, "*.docx")) + glob.glob(os.path.join(input_dir, "*.rtf"))

# Dashboard Stats
col1, col2, col3 = st.columns(3)
with col1:
    st.markdown(f"""<div class="stat-card"><div class="stat-value">{len(docx_files)}</div><div class="stat-label">Documents Found</div></div>""", unsafe_allow_html=True)
with col2:
    st.markdown(f"""<div class="stat-card"><div class="stat-value">{sum([rename_option, remove_colors_option, convert_md_option])}</div><div class="stat-label">Active Operations</div></div>""", unsafe_allow_html=True)
with col3:
    st.markdown(f"""<div class="stat-card"><div class="stat-value">Ready</div><div class="stat-label">System Status</div></div>""", unsafe_allow_html=True)

st.markdown("---")

if start_btn:
    if not docx_files:
        st.error("No files found to process!")
    else:
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            st.toast(f"Created output directory: {output_dir}")

        progress_bar = st.progress(0)
        status_text = st.empty()

        results_data = []

        start_time = time.time()

        for i, file_path in enumerate(docx_files):
            filename = os.path.basename(file_path)
            status_text.markdown(f"**Processing:** `{filename}`")

            file_result = {
                "Original File": filename,
                "New Name": "-",
                "Cleaned": "Skipped",
                "Markdown": "Skipped",
                "Status": "Success"
            }

            try:
                # 1. Rename Logic
                if rename_option:
                    name, ext = os.path.splitext(filename)
                    new_name = to_snake_case(name)
                    new_filename = f"{new_name}{ext}"
                    file_result["New Name"] = new_filename
                else:
                    new_filename = filename
                    file_result["New Name"] = filename

                # Copy to output
                temp_output_path = os.path.join(output_dir, new_filename)
                shutil.copy2(file_path, temp_output_path)

                # 2. Remove Colors
                if remove_colors_option and temp_output_path.lower().endswith('.docx'):
                    clean_path = temp_output_path.replace('.docx', '_clean.docx')
                    success, msg = remove_colors(temp_output_path, clean_path)
                    if success:
                        # If successful, we might want to use the clean file for MD conversion
                        # Or keep both. Let's keep the clean file as the "main" one for conversion if desired.
                        # For now, let's just mark it done.
                        file_result["Cleaned"] = "Done"
                        # Update temp_output_path to point to clean file if we want MD conversion on clean version
                        temp_output_path = clean_path
                    else:
                        file_result["Cleaned"] = "Failed"
                        file_result["Status"] = "Partial Error"
                        st.error(f"Color removal failed for {filename}: {msg}")

                # 3. Convert to MD
                if convert_md_option:
                    success, output_md = convert_docx_to_md(temp_output_path, output_dir)
                    if success:
                        file_result["Markdown"] = os.path.basename(output_md)
                    else:
                        file_result["Markdown"] = "Failed"
                        file_result["Status"] = "Partial Error"
                        st.error(f"MD Conversion failed for {filename}: {msg}")

            except Exception as e:
                file_result["Status"] = "Error"
                st.error(f"Error processing {filename}: {str(e)}")

            results_data.append(file_result)
            progress_bar.progress((i + 1) / len(docx_files))

        end_time = time.time()
        duration = end_time - start_time

        progress_bar.empty()
        status_text.empty()

        st.success(f"🎉 Processing complete in {duration:.2f} seconds!")

        # Results Table
        df = pd.DataFrame(results_data)
        st.dataframe(
            df,
            column_config={
                "Status": st.column_config.TextColumn(
                    "Status",
                    help="Processing Status",
                    validate="^Success$"
                ),
            },
            use_container_width=True,
            hide_index=True
        )

        # Download Report (Optional, simple CSV)
        csv = df.to_csv(index=False).encode('utf-8')
        st.download_button(
            label="📥 Download Processing Report",
            data=csv,
            file_name='processing_report.csv',
            mime='text/csv',
        )

# Footer
st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #6c757d; font-size: 0.8rem;">
    DOCX Processor Pro | Built with Streamlit
</div>
""", unsafe_allow_html=True)

