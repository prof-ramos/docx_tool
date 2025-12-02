"""
Cleaner module for removing formatting from DOCX files.
"""
from pathlib import Path
from typing import Any, Optional

import docx
from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from docx.text.run import Run

class Cleaner:
    """Handles document cleaning operations."""

    def remove_colors(self, input_path: Path, output_path: Path) -> None:
        """
        Removes colors and highlighting from a DOCX file.

        Args:
            input_path: Path to the input DOCX file.
            output_path: Path to save the cleaned DOCX file.

        Raises:
            Exception: If file processing fails.
        """
        doc = docx.Document(input_path)
        self._clean_document(doc)
        doc.save(output_path)

    def _clean_document(self, doc: Document) -> None:
        """
        Internal method to clean the document structure.

        Args:
            doc: The python-docx Document object.
        """
        # Process paragraphs
        for paragraph in doc.paragraphs:
            self._clear_shading(paragraph)
            for run in paragraph.runs:
                self._clean_run(run)

        # Process tables
        for table in doc.tables:
            self._clean_table(table)

    def _clean_table(self, table: Table) -> None:
        """Cleans a table structure."""
        for row in table.rows:
            for cell in row.cells:
                self._clear_shading(cell)
                for paragraph in cell.paragraphs:
                    self._clear_shading(paragraph)
                    for run in paragraph.runs:
                        self._clean_run(run)

    def _clean_run(self, run: Run) -> None:
        """
        Clears formatting from a run.

        Args:
            run: The run object to clean.
        """
        run.font.color.rgb = None
        if run.font.highlight_color is not None:
            run.font.highlight_color = None

        # Access XML element for deeper cleaning
        rPr = run._element.get_or_add_rPr()

        # Remove w:highlight
        for highlight in rPr.xpath('./w:highlight'):
            highlight.getparent().remove(highlight)

        # Remove w:shd in run (background shading)
        for shd in rPr.xpath('./w:shd'):
            shd.getparent().remove(shd)

    def _clear_shading(self, element: Any) -> None:
        """
        Clears shading from an element (paragraph or table cell).

        Args:
            element: The element (Paragraph or _Cell) to clean.
        """
        # Check if element has access to underlying XML element
        if hasattr(element, '_element'):
            xml_element = element._element
            if xml_element.xpath('.//w:shd'):
                for shd in xml_element.xpath('.//w:shd'):
                    shd.getparent().remove(shd)
