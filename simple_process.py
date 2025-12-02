#!/usr/bin/env python3
"""Simple DOCX processor without heavy dependencies."""
import os
import sys
from pathlib import Path
from docx import Document
import unicodedata
import re

def to_snake_case(name):
    """Convert string to snake_case."""
    name = unicodedata.normalize('NFKD', name).encode('ASCII', 'ignore').decode('ASCII')
    name = name.lower()
    name = re.sub(r'[^a-z0-9]', '_', name)
    name = re.sub(r'_+', '_', name)
    return name.strip('_')

def remove_colors(input_path, output_path):
    """Remove colors and highlighting from DOCX."""
    try:
        doc = Document(input_path)

        def clear_shading(element):
            if element._element.xpath('.//w:shd'):
                for shd in element._element.xpath('.//w:shd'):
                    shd.getparent().remove(shd)

        # Process paragraphs
        for paragraph in doc.paragraphs:
            clear_shading(paragraph)
            for run in paragraph.runs:
                run.font.color.rgb = None
                if run.font.highlight_color is not None:
                    run.font.highlight_color = None

                rPr = run._element.get_or_add_rPr()
                if rPr.xpath('./w:highlight'):
                    for highlight in rPr.xpath('./w:highlight'):
                        highlight.getparent().remove(highlight)
                if rPr.xpath('./w:shd'):
                    for shd in rPr.xpath('./w:shd'):
                        shd.getparent().remove(shd)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    clear_shading(cell)
                    for paragraph in cell.paragraphs:
                        clear_shading(paragraph)
                        for run in paragraph.runs:
                            run.font.color.rgb = None
                            if run.font.highlight_color is not None:
                                run.font.highlight_color = None

                            rPr = run._element.get_or_add_rPr()
                            if rPr.xpath('./w:highlight'):
                                for highlight in rPr.xpath('./w:highlight'):
                                    highlight.getparent().remove(highlight)
                            if rPr.xpath('./w:shd'):
                                for shd in rPr.xpath('./w:shd'):
                                    shd.getparent().remove(shd)

        doc.save(output_path)
        return True, "Colors removed successfully"
    except Exception as e:
        return False, str(e)

def process_directory(input_dir, output_dir, clean=False):
    """Process all DOCX files in directory."""
    input_path = Path(input_dir)
    output_path = Path(output_dir)

    if not input_path.exists():
        print(f"❌ Input directory not found: {input_dir}")
        return

    # Create output directory
    output_path.mkdir(parents=True, exist_ok=True)

    # Find DOCX files
    docx_files = list(input_path.glob("*.docx"))

    if not docx_files:
        print(f"⚠️ No DOCX files found in {input_dir}")
        return

    print(f"📄 Found {len(docx_files)} files")
    print()

    success_count = 0
    error_count = 0

    for file_path in docx_files:
        print(f"Processing: {file_path.name}")

        try:
            # Generate output filename
            name_stem = file_path.stem
            normalized_name = to_snake_case(name_stem)

            if clean:
                output_file = output_path / f"{normalized_name}_clean.docx"
                success, msg = remove_colors(file_path, output_file)

                if success:
                    print(f"  ✅ Cleaned: {output_file.name}")
                    success_count += 1
                else:
                    print(f"  ❌ Error: {msg}")
                    error_count += 1
            else:
                # Just copy with normalized name
                output_file = output_path / f"{normalized_name}.docx"
                doc = Document(file_path)
                doc.save(output_file)
                print(f"  ✅ Copied: {output_file.name}")
                success_count += 1

        except Exception as e:
            print(f"  ❌ Error: {str(e)}")
            error_count += 1

    print()
    print("="* 50)
    print(f"✅ Success: {success_count}")
    print(f"❌ Errors: {error_count}")
    print(f"📁 Output: {output_dir}")

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Simple DOCX processor")
    parser.add_argument("input_dir", help="Input directory")
    parser.add_argument("--output-dir", "-o", default="Output", help="Output directory")
    parser.add_argument("--clean", action="store_true", help="Remove colors")

    args = parser.parse_args()

    process_directory(args.input_dir, args.output_dir, args.clean)
