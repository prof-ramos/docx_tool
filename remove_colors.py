import docx
from docx.shared import RGBColor
from docx.oxml.ns import qn

def remove_colors(input_path, output_path):
    doc = docx.Document(input_path)

    # Function to clear shading from an element (paragraph or table cell)
    def clear_shading(element):
        if element._element.xpath('.//w:shd'):
            for shd in element._element.xpath('.//w:shd'):
                shd.getparent().remove(shd)

    # Process paragraphs
    for paragraph in doc.paragraphs:
        clear_shading(paragraph)
        for run in paragraph.runs:
            run.font.color.rgb = None
            # Clear highlight
            if run.font.highlight_color is not None:
                run.font.highlight_color = None
            # Also need to check for w:highlight tag directly as python-docx might not catch all
            rPr = run._element.get_or_add_rPr()
            if rPr.xpath('./w:highlight'):
                for highlight in rPr.xpath('./w:highlight'):
                    highlight.getparent().remove(highlight)

            # Clear background shading in runs (less common but possible)
            if rPr.xpath('./w:shd'):
                for shd in rPr.xpath('./w:shd'):
                    shd.getparent().remove(shd)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                clear_shading(cell)
                for paragraph in cell.paragraphs:
                    clear_shading(paragraph)
                    for run in paragraph.runs:
                        run.font.color.rgb = None
                        if run.font.highlight_color is not None:
                            run.font.highlight_color = None

                        rPr = run._element.get_or_add_rPr()
                        if rPr.xpath('./w:highlight'):
                            for highlight in rPr.xpath('./w:highlight'):
                                highlight.getparent().remove(highlight)
                        if rPr.xpath('./w:shd'):
                            for shd in rPr.xpath('./w:shd'):
                                shd.getparent().remove(shd)

    doc.save(output_path)
    print(f"Processed file saved to: {output_path}")

if __name__ == "__main__":
    input_file = "Processo Administrativo (Lei 9784).docx"
    output_file = "Processo Administrativo (Lei 9784)_clean.docx"
    remove_colors(input_file, output_file)
